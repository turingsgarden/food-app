"""
OCR health-metric extraction pipeline for food-app.

Design implemented:
1. Transcribe the uploaded image/PDF without rewriting the text.
2. Extract 11 canonical health metrics including additional tests while preserving raw label/value/unit.
3. Normalize units.
4. Validate values against approved ranges.
5. Return one of: ok, no_fields, no_text.
6. Preserve a field-level raw + processed + transform structure.

Required environment variables:
    GEMINI_API_KEY
Optional:
    GEMINI_OCR_MODEL: gemini-3.1-flash-lite
"""

from __future__ import annotations

import json
import math
import mimetypes
import os
import re
import time
import io
import docx
from dataclasses import dataclass
from typing import Any, Dict, Iterable, Mapping, Optional, Tuple

import fitz  # PyMuPDF

try:
    from google import genai
    from google.genai import errors, types
    _GENAI_IMPORT_ERROR = None
except ImportError as import_error:  # Allows non-API unit tests to run.
    genai = None
    errors = None
    types = None
    _GENAI_IMPORT_ERROR = import_error


FIELD_NAMES = (
    "systolic_bp",
    "diastolic_bp",
    "height_cm",
    "weight_kg",
    "bmi",
    "blood_sugar",
    "hba1c",
    "cholesterol",
    "ldl",
    "hdl",
    "triglycerides",
)


@dataclass(frozen=True)
class FieldSpec:
    value_type: str
    standard_unit: str
    minimum: float
    maximum: float
    aliases: Tuple[str, ...]

FIELD_SPECS: Dict[str, FieldSpec] = {
    "systolic_bp": FieldSpec("int", "mmHg", 60, 250, ("systolic blood pressure", "systolic pressure", "systolic", "SBP", "high pressure", "left value of BP")),
    "diastolic_bp": FieldSpec("int", "mmHg", 30, 150, ("diastolic blood pressure", "diastolic pressure", "diastolic", "DBP", "low pressure", "right value of BP")),
    "height_cm": FieldSpec("float", "cm", 100, 230, ("body height", "height", "Ht")),
    "weight_kg": FieldSpec("float", "kg", 30, 300, ("body weight", "weight", "Wt", "BW")),
    "bmi": FieldSpec("float", "kg/m²", 10, 60, ("body mass index", "BMI")),
    "blood_sugar": FieldSpec("float", "mmol/L", 2.0, 30.0, ("fasting blood glucose", "fasting plasma glucose", "fasting glucose", "blood glucose", "blood sugar", "glucose", "FBG", "FPG", "GLU")),
    "hba1c": FieldSpec("float", "%", 3.0, 20.0, ("glycated hemoglobin", "glycosylated hemoglobin", "HbA1c", "A1c")),
    "cholesterol": FieldSpec("float", "mmol/L", 1.0, 15.0, ("total cholesterol", "cholesterol total", "total chol", "TC", "CHOL")),
    "ldl": FieldSpec("float", "mmol/L", 0.3, 10.0, ("low-density lipoprotein cholesterol", "low density lipoprotein", "LDL-C", "LDL")),
    "hdl": FieldSpec("float", "mmol/L", 0.3, 5.0, ("high-density lipoprotein cholesterol", "high density lipoprotein", "HDL-C", "HDL")),
    "triglycerides": FieldSpec("float", "mmol/L", 0.2, 10.0, ("triglycerides", "triglyceride", "TRIG", "TG")),
}

DOCX_MIME_TYPE = (
    "application/vnd.openxmlformats-officedocument."
    "wordprocessingml.document"
)
SUPPORTED_MIME_TYPES = {
    "image/jpeg",
    "image/png",
    "image/webp",
    "image/heic",
    "image/heif",
    "application/pdf",
    DOCX_MIME_TYPE,
}


MAX_PDF_PAGES = 10
MAX_TRANSCRIPTION_CHARS = 80_000

# for 11 main health fields
RAW_ITEM_SCHEMA = {
    "anyOf": [
        {"type": "null"},
        {
            "type": "object",
            "properties": {
                "name": {"anyOf": [{"type": "string"}, {"type": "null"}]},
                "value": {
                    "anyOf": [
                        {"type": "string"},
                        {"type": "number"},
                        {"type": "null"},
                    ]
                },
                "unit": {"anyOf": [{"type": "string"}, {"type": "null"}]},
            },
            "required": ["name", "value", "unit"],
            "additionalProperties": False,
        },
    ]
}

#extra health fields
ADDITIONAL_ITEM_SCHEMA = {
    "type": "object",
    "properties": {
        "name": {"type": "string"},
        "value": {
            "anyOf": [
                {"type": "string"},
                {"type": "number"},
                {"type": "null"},
            ]
        },
        "unit": {"anyOf": [{"type": "string"}, {"type": "null"},]},
    },
    "required": ["name", "value", "unit"],
    "additionalProperties": False,
}

#The complete Gemini extraction response
RAW_EXTRACTION_SCHEMA = {
    "type": "object",
    "properties": {
        **{name: RAW_ITEM_SCHEMA for name in FIELD_NAMES},

        "additional_fields": {
            "type": "array",
            "items": ADDITIONAL_ITEM_SCHEMA,
        },
    },
    "required": [*FIELD_NAMES, "additional_fields"],
    "additionalProperties": False,
}



class HealthOcrError(RuntimeError):
    """Base exception for OCR feature failures."""


class HealthOcrConfigurationError(HealthOcrError):
    """Raised when required server configuration is missing."""


class HealthOcrApiError(HealthOcrError):
    """Raised when Gemini cannot complete an OCR request."""


def _require_genai_sdk() -> None:
    if genai is None or errors is None or types is None:
        raise HealthOcrConfigurationError(
            "The google-genai package is not installed. "
            "Run: pip install google-genai"
        ) from _GENAI_IMPORT_ERROR

#Reads GEMINI_API_KEY from the backend environment and creates a Gemini API client.
def _get_client() -> "genai.Client":
    _require_genai_sdk()
    api_key = os.getenv("GEMINI_API_KEY", "").strip()
    if not api_key:
        raise HealthOcrConfigurationError(
            "GEMINI_API_KEY is not configured on the backend."
        )
    return genai.Client(api_key=api_key)


def _model_name() -> str:
    return os.getenv("GEMINI_OCR_MODEL", "gemini-3.5-flash").strip()


def _guess_mime_type(filename: str, supplied_mime_type: Optional[str]) -> str:
    mime_type = (supplied_mime_type or "").split(";")[0].strip().lower()

    if mime_type in {"image/jpg", "image/pjpeg"}:
        mime_type = "image/jpeg"

    if not mime_type or mime_type == "application/octet-stream":
        guessed, _ = mimetypes.guess_type(filename or "")
        mime_type = (guessed or "").lower()

    extension = os.path.splitext(filename or "")[1].lower()
    extension_fallbacks = {
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".png": "image/png",
        ".webp": "image/webp",
        ".heic": "image/heic",
        ".heif": "image/heif",
        ".pdf": "application/pdf",
        ".docx": DOCX_MIME_TYPE,
    }
    if mime_type not in SUPPORTED_MIME_TYPES:
        mime_type = extension_fallbacks.get(extension, mime_type)

    if mime_type not in SUPPORTED_MIME_TYPES:
        raise ValueError(
            f"Unsupported document type: {mime_type or 'unknown'}. "
            "Use JPEG, PNG, WebP, HEIC, HEIF, PDF, or DOCX."
        )

    return mime_type


def _retry_generate_content(
    *,
    contents: Iterable[Any],
    config: types.GenerateContentConfig,
    max_retries: int = 3,
):
    client = _get_client()
    last_error: Optional[Exception] = None
    retryable_codes = {408, 429, 500, 502, 503, 504}

    for attempt in range(1, max_retries + 1):
        try:
            return client.models.generate_content(
                model=_model_name(),
                contents=list(contents),
                config=config,
            )
        except errors.APIError as error:
            last_error = error
            status_code = getattr(error, "code", None)
            if attempt < max_retries and status_code in retryable_codes:
                time.sleep(min(2 ** attempt, 8))
                continue
            raise HealthOcrApiError(
                getattr(error, "message", None) or str(error)
            ) from error
        except Exception as error:
            last_error = error
            if attempt < max_retries:
                time.sleep(min(2 ** attempt, 8))
                continue
            raise HealthOcrApiError(str(error)) from error

    raise HealthOcrApiError(str(last_error or "Unknown Gemini API error"))


def _gemini_transcribe_image(image_bytes: bytes, mime_type: str, *, max_retries: int = 3,) -> str:

    """Transcribe one image verbatim. Do not normalize names or units here."""

    _require_genai_sdk()
    prompt = """
Transcribe every readable character from this medical or laboratory report.

Rules:
- Preserve the original wording, abbreviations, numbers, symbols, and units.
- Preserve line order as much as possible.
- Do not rename tests.
- Do not convert units.
- Do not calculate values.
- Do not summarize or explain.
- Return only the transcription.
- If no text is readable, return an empty string.
""".strip()

    response = _retry_generate_content(
        contents=[
            prompt,
            types.Part.from_bytes(data=image_bytes, mime_type=mime_type),
        ],
        config=types.GenerateContentConfig(
            temperature=0,
            max_output_tokens=8192,
            http_options=types.HttpOptions(timeout=60_000),
        ),
        max_retries=max_retries,
    )
    return (response.text or "").strip()


def _extract_pdf_text_or_images(pdf_bytes: bytes,) -> Tuple[str, Tuple[bytes, ...]]:
    """
    Return embedded text when the PDF has a usable text layer.
    Otherwise render pages to PNG for image transcription.
    """
    try:
        document = fitz.open(stream=pdf_bytes, filetype="pdf")
    except Exception as error:
        raise ValueError(f"Invalid or unreadable PDF: {error}") from error

    try:
        page_count = min(document.page_count, MAX_PDF_PAGES)
        embedded_parts = []

        for page_index in range(page_count):
            page_text = document.load_page(page_index).get_text("text").strip()
            if page_text:
                embedded_parts.append(page_text)

        embedded_text = "\n\n".join(embedded_parts).strip()

        # Use the text layer when it contains enough actual content.
        if len(re.sub(r"\s+", "", embedded_text)) >= 20:
            return embedded_text[:MAX_TRANSCRIPTION_CHARS], tuple()

        rendered_pages = []
        matrix = fitz.Matrix(2.0, 2.0)

        for page_index in range(page_count):
            page = document.load_page(page_index)
            pixmap = page.get_pixmap(matrix=matrix, alpha=False)
            rendered_pages.append(pixmap.tobytes("png"))

        return "", tuple(rendered_pages)
    finally:
        document.close()

def _extract_docx_text(docx_bytes: bytes) -> str:
    """
    Extract readable text from DOCX paragraphs and tables.
    The extracted text is then passed through the same new
    health-field extraction pipeline used for photos and PDFs.
    """
    try:
        document = docx.Document(io.BytesIO(docx_bytes))
    except Exception as error:
        raise ValueError(
            f"Invalid or unreadable DOCX file: {error}"
        ) from error

    text_parts: list[str] = []

    # Normal paragraphs
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            text_parts.append(text)

    # Medical results are frequently inside Word tables.
    for table in document.tables:
        for row in table.rows:
            cells = [
                re.sub(r"\s+", " ", cell.text).strip()
                for cell in row.cells
            ]

            if any(cells):
                text_parts.append("\t".join(cells))

    return "\n".join(text_parts).strip()[
        :MAX_TRANSCRIPTION_CHARS
    ]

def transcribe_document(
    file_bytes: bytes,
    *,
    filename: str = "upload.jpg",
    mime_type: Optional[str] = None,
    max_retries: int = 3,
) -> str:
    if not file_bytes:
        return ""

    resolved_mime = _guess_mime_type(
        filename,
        mime_type,
    )

    # DOCX: extract text from paragraphs and tables.
    if resolved_mime == DOCX_MIME_TYPE:
        return _extract_docx_text(file_bytes)

    # Images: send directly to Gemini transcription.
    if resolved_mime != "application/pdf":
        return _gemini_transcribe_image(
            file_bytes,
            resolved_mime,
            max_retries=max_retries,
        )[:MAX_TRANSCRIPTION_CHARS]

    # PDF: use embedded text when available.
    # Otherwise render its pages and run Gemini OCR.
    embedded_text, page_images = (
        _extract_pdf_text_or_images(file_bytes)
    )

    if embedded_text:
        return embedded_text

    page_texts = []

    for page_number, page_bytes in enumerate(
        page_images,
        start=1,
    ):
        page_text = _gemini_transcribe_image(
            page_bytes,
            "image/png",
            max_retries=max_retries,
        )

        if page_text:
            page_texts.append(
                f"[Page {page_number}]\n{page_text}"
            )

    return "\n\n".join(page_texts).strip()[
        :MAX_TRANSCRIPTION_CHARS
    ]

def _extraction_prompt(transcription: str) -> str:
    alias_lines = []

    for field_name, spec in FIELD_SPECS.items():
        alias_lines.append(
            f"- {field_name}: {', '.join(spec.aliases)}"
        )

    return f"""
Extract the following 11 canonical health metrics AND every other laboratory
or health measurement found in the report transcription.

Canonical fields and accepted label examples:
{chr(10).join(alias_lines)}

For the 11 canonical fields:
- Put each matched measurement in its corresponding canonical JSON key.
- Return null when the field is absent.

For every other measurement:
- Put it inside `additional_fields`.
- Return a short, normalized display name, preferably 2–5 words.
- Do not return the complete report description as the test name.
- Remove score ranges, measurement descriptions, report status, brackets,
  codes, specimen details and method information from test names.
- Use common medical abbreviations when they are clearer and shorter.
- Examples:
  "Pain severity - 0-10 verbal numeric rating [Score] - Reported"
  -> "Pain Severity"
  "Patient Health Questionnaire-9: Modified for Teens total score"
  -> "PHQ-9"
  "White Blood Cell Count"
  -> "WBC Count"
  "Respiratory rate"
  -> "Respiratory Rate"
  - Do not repeat the unit in the test name.
- When the unit is "score", omit "Score" and "Total Score" from the test name.
- Preserve the original result.
- Preserve the original unit.
- One measurement must produce one object.

Important rules:
- Matching is case-insensitive.
- Preserve labels, values, and units exactly as they appear.
- Do not invent, interpret, diagnose, or calculate values.
- Do not duplicate one of the 11 canonical fields inside `additional_fields`.
- Do not include patient name, patient ID, address, age, gender, dates,
  doctor names, signatures, headings, reference ranges, or report metadata.
- Include text results such as Positive, Negative, Reactive,
  Non-reactive, Detected, or Not Detected.
- Do not place LDL or HDL values into cholesterol.
- cholesterol means Total Cholesterol / Total / TC only.
- SBP and DBP must never be swapped.
- Return every required JSON key.
- Return only JSON matching the supplied schema.

REPORT TRANSCRIPTION:
--- BEGIN REPORT ---
{transcription}
--- END REPORT ---
""".strip()

ADDITIONAL_DISPLAY_NAME_PATTERNS = (
    (
        r"\bpain severity\b",
        "Pain Severity",
    ),
    (
        r"\bpatient health questionnaire(?:\s+9)?\b|\bphq\s+9\b",
        "PHQ-9",
    ),
    (
        r"\bheart rate\b|\bpulse rate\b",
        "Heart Rate",
    ),
    (
        r"\brespiratory rate\b|\brespiration rate\b",
        "Respiratory Rate",
    ),
    (
        r"\bwhite blood cell(?: count)?\b|\bwbc(?: count)?\b",
        "WBC Count",
    ),
    (
        r"\bred blood cell(?: count)?\b|\brbc(?: count)?\b",
        "RBC Count",
    ),
    (
        r"\bplatelet(?: count)?\b|\bplt(?: count)?\b",
        "Platelet Count",
    ),
)
def _normalize_additional_test_name(name: str) -> str:
    original_key = _normalized_label(name)

    # Handle long descriptive names using partial/pattern matching.
    for pattern, display_name in ADDITIONAL_DISPLAY_NAME_PATTERNS:
        if re.search(pattern, original_key):
            return display_name

    # Remove content inside square or curly brackets.
    cleaned_name = re.sub(
        r"\[[^\]]*\]|\{[^}]*\}",
        " ",
        name,
    )

    # Remove descriptions following a spaced dash.
    cleaned_name = re.split(
        r"\s+(?:-|–|—)\s+",
        cleaned_name,
        maxsplit=1,
    )[0]

    cleaned_name = re.sub(
        r"\s+",
        " ",
        cleaned_name,
    ).strip(" :-")

    lookup_key = _normalized_label(cleaned_name)

    # Existing exact alias lookup.
    normalized_name = _ADDITIONAL_TEST_LOOKUP.get(lookup_key)

    if normalized_name:
        return normalized_name

    return cleaned_name


def _clean_json_text(text: str) -> str:
    cleaned = (text or "").strip()
    cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\s*```$", "", cleaned)
    return cleaned.strip()


def _empty_raw_item() -> Dict[str, None]:
    return {"name": None, "value": None, "unit": None}


def _empty_raw_extraction() -> Dict[str, Dict[str, None]]:
    return {name: _empty_raw_item() for name in FIELD_NAMES}


def _validate_raw_extraction(data: Any) -> Dict[str, Dict[str, Any]]:
    normalized = _empty_raw_extraction()

    if not isinstance(data, Mapping):
        return normalized

    for field_name in FIELD_NAMES:
        item = data.get(field_name)

        if item is None:
            continue

        if not isinstance(item, Mapping):
            continue

        raw_name = item.get("name")
        raw_value = item.get("value")
        raw_unit = item.get("unit")

        normalized[field_name] = {
            "name": None if raw_name is None else str(raw_name).strip() or None,
            "value": raw_value,
            "unit": None if raw_unit is None else str(raw_unit).strip() or None,
        }

    return normalized

######### for additional fields #################

def _normalized_label(value: Any) -> str:
    text = "" if value is None else str(value)

    return re.sub(
        r"[^a-z0-9]+",
        " ",
        text.lower(),
    ).strip()


def _is_canonical_label(
    name: str,
    canonical_raw: Optional[
        Mapping[str, Mapping[str, Any]]
    ] = None,
) -> bool:
    candidate = _normalized_label(name)

    if not candidate:
        return False

    # Remove exact duplicates of canonical fields selected by Gemini.
    if canonical_raw:
        for field_name in FIELD_NAMES:
            item = canonical_raw.get(field_name) or {}
            raw_name = _normalized_label(item.get("name"))

            if raw_name and candidate == raw_name:
                return True

    # Also compare against configured aliases.
    aliases = set(FIELD_NAMES)

    for spec in FIELD_SPECS.values():
        aliases.update(spec.aliases)

    padded_candidate = f" {candidate} "

    for alias in aliases:
        alias_key = _normalized_label(alias)

        if not alias_key:
            continue

        if candidate == alias_key:
            return True

        # Longer aliases may appear inside a complete laboratory label.
        if (
            len(alias_key) >= 3
            and f" {alias_key} " in padded_candidate
        ):
            return True

    return False


def _is_metadata_label(name: str) -> bool:
    candidate = _normalized_label(name)

    metadata_labels = {
        "patient name",
        "name",
        "age",
        "gender",
        "sex",
        "date of birth",
        "dob",
        "address",
        "city",
        "state",
        "postcode",
        "postal code",
        "mrn",
        "patient id",
        "report date",
        "collection date",
        "received date",
        "doctor",
        "physician",
        "laboratory",
        "lab name",
        "signature",
    }

    return candidate in metadata_labels


def _validate_additional_fields(
    data: Any,
    canonical_raw: Optional[
        Mapping[str, Mapping[str, Any]]
    ] = None,
) -> list[Dict[str, Any]]:
    if not isinstance(data, Mapping):
        return []

    raw_items = data.get("additional_fields")

    if not isinstance(raw_items, list):
        return []

    normalized: list[Dict[str, Any]] = []
    seen = set()

    for item in raw_items[:200]:
        if not isinstance(item, Mapping):
            continue

        name = str(item.get("name") or "").strip()
        value = item.get("value")
        unit = str(item.get("unit") or "").strip() or None
        if unit:
            unit = re.sub(r"[{}\[\]()]", "", unit,).strip() or None   


        # Avoid repeating "score" in both the name and unit.
        # if unit and _normalized_label(unit) == "score":
        #     name = re.sub(
        #         r"\s+(?:total\s+)?score\s*$",
        #         "",
        #         name,
        #         flags=re.IGNORECASE,
        #     ).strip()

        if (
            not name
            or _is_canonical_label(name, canonical_raw)
            or _is_metadata_label(name)
        ):
            continue

        if isinstance(value, str):
            value = value.strip()

        dedupe_key = (
            _normalized_label(name),
            _normalized_label(value),
            _normalized_label(unit),
        )

        if dedupe_key in seen:
            continue

        seen.add(dedupe_key)

        normalized.append({
            "name": name,
            "value": value,
            "unit": unit,
        })

    return normalized

#_______
def _extract_health_values_from_text(
    transcription: str,
    *,
    max_retries: int = 3,
) -> Tuple[
    Dict[str, Dict[str, Any]],
    list[Dict[str, Any]],
]:
    """
    Extract the 11 canonical fields and all other medical measurements.
    """
    _require_genai_sdk()

    response = _retry_generate_content(
        contents=[_extraction_prompt(transcription)],
        config=types.GenerateContentConfig(
            temperature=0,
            max_output_tokens=8192,
            response_mime_type="application/json",
            response_json_schema=RAW_EXTRACTION_SCHEMA,
            http_options=types.HttpOptions(
                timeout=120_000
            ),
        ),
        max_retries=max_retries,
    )

    parsed = getattr(response, "parsed", None)

    if not isinstance(parsed, Mapping):
        raw_text = _clean_json_text(
            response.text or ""
        )
        parsed = json.loads(raw_text)

    canonical_raw = _validate_raw_extraction(parsed)

    additional_fields = _validate_additional_fields(
        parsed,
        canonical_raw=canonical_raw,
    )

    return canonical_raw, additional_fields


NUMBER_PATTERN = r"[-+]?\d+(?:[.,]\d+)?"
UNIT_PATTERN = (
    r"(?:mm\s*(?:\[\s*hg\s*\]|\(\s*hg\s*\)|hg)|"
    r"kpa|mg\s*/?\s*dl|mmol\s*/\s*l|mmol\s*/\s*mol|"
    r"kg\s*/\s*m(?:2|²)|%|cm|mm|m|inches?|inch|in|\"|kgs?|"
    r"kilograms?|lbs?|pounds?)"
)


def _alias_regex(alias: str) -> str:
    escaped = re.escape(alias)
    escaped = escaped.replace(r"\ ", r"\s*")

    # Latin abbreviations should not match inside larger words.
    if re.fullmatch(r"[A-Za-z0-9\-\s]+", alias):
        return rf"(?<![A-Za-z0-9]){escaped}(?![A-Za-z0-9])"

    return escaped


def _find_line_value(
    line: str,
    aliases: Tuple[str, ...],
) -> Optional[Dict[str, Any]]:
    alias_group = "|".join(
        _alias_regex(alias)
        for alias in sorted(aliases, key=len, reverse=True)
    )

    patterns = (
        # Label [unit] : value
        rf"(?P<name>{alias_group})\s*[\(\[]?\s*(?P<unit1>{UNIT_PATTERN})?"
        rf"\s*[\)\]]?\s*[:=\-–—]?\s*(?P<value>{NUMBER_PATTERN})",
        # Label : value unit
        rf"(?P<name>{alias_group})\s*[:=\-–—]?\s*(?P<value>{NUMBER_PATTERN})"
        rf"\s*(?P<unit2>{UNIT_PATTERN})?",
        # Label ... value ... unit, allowing a small amount of table noise
        rf"(?P<name>{alias_group}).{{0,24}}?(?P<value>{NUMBER_PATTERN})"
        rf"\s*(?P<unit2>{UNIT_PATTERN})?",
    )

    for pattern in patterns:
        match = re.search(pattern, line, flags=re.IGNORECASE)
        if not match:
            continue

        unit = match.groupdict().get("unit1") or match.groupdict().get("unit2")
        return {
            "name": match.group("name").strip(),
            "value": match.group("value").strip(),
            "unit": unit.strip() if unit else None,
        }

    return None


def _extract_health_values_regex(
    transcription: str,
) -> Dict[str, Dict[str, Any]]:
    """
    It is a backup extractor. It runs when Gemini cannot extract the information properly.
    It is intentionally conservative and does not guess unrelated tests.
    """
    result = _empty_raw_extraction()
    lines = [line.strip() for line in transcription.splitlines() if line.strip()]

    # Combined blood-pressure notation, for example BP 120/80 mmHg.
    bp_pattern = re.compile(
        rf"(?:(?:blood\s*pressure|\bBP\b)\s*[:=\-]?\s*)?"
        rf"(?P<sys>\d{{2,3}}(?:\.\d+)?)\s*/\s*"
        rf"(?P<dia>\d{{2,3}}(?:\.\d+)?)"
        rf"\s*(?P<unit>mm\s*(?:\[\s*hg\s*\]|\(\s*hg\s*\)|hg)|kpa)?"
        rf"(?!\s*/\s*\d)",
        flags=re.IGNORECASE,
    )

    for line in lines:
        lower_line = line.lower()

        # Do not interpret obvious date lines as blood pressure.
        if not any(
            token in lower_line
            for token in ("bp", "blood pressure", "systolic", "diastolic", "收缩压", "舒张压")
        ):
            continue

        match = bp_pattern.search(line)
        if match:
            unit = match.group("unit") or "mmHg"
            result["systolic_bp"] = {
                "name": "SBP",
                "value": match.group("sys"),
                "unit": unit,
            }
            result["diastolic_bp"] = {
                "name": "DBP",
                "value": match.group("dia"),
                "unit": unit,
            }
            break

    for field_name, spec in FIELD_SPECS.items():
        if result[field_name]["value"] is not None:
            continue

        for line in lines:
            # Total cholesterol must never be read from an LDL or HDL row.
            if field_name == "cholesterol":
                lower = line.lower()
                if "ldl" in lower or "hdl" in lower:
                    continue

            found = _find_line_value(line, spec.aliases)
            if found is not None:
                result[field_name] = found
                break

    return result


def _merge_raw_extractions(
    primary: Mapping[str, Mapping[str, Any]],
    fallback: Mapping[str, Mapping[str, Any]],
) -> Dict[str, Dict[str, Any]]:
    merged = _empty_raw_extraction()

    for field_name in FIELD_NAMES:
        primary_item = primary.get(field_name) or {}
        fallback_item = fallback.get(field_name) or {}

        if primary_item.get("value") is not None:
            chosen = primary_item
        else:
            chosen = fallback_item

        merged[field_name] = {
            "name": chosen.get("name"),
            "value": chosen.get("value"),
            "unit": chosen.get("unit"),
        }

    return merged


def _coerce_number(value: Any) -> Optional[float]:
    if value is None or isinstance(value, bool):
        return None

    if isinstance(value, (int, float)):
        numeric = float(value)
        return numeric if math.isfinite(numeric) else None

    text = str(value).strip()
    if not text:
        return None

    # Remove thousands separators but preserve a decimal comma.
    text = text.replace("−", "-")
    if re.fullmatch(r"[-+]?\d{1,3}(?:,\d{3})+(?:\.\d+)?", text):
        text = text.replace(",", "")
    else:
        text = text.replace(",", ".")

    match = re.search(NUMBER_PATTERN, text)
    if not match:
        return None

    try:
        numeric = float(match.group(0).replace(",", "."))
    except ValueError:
        return None

    return numeric if math.isfinite(numeric) else None


def _normalize_unit_text(unit: Optional[str]) -> str:
    text = (unit or "").strip().lower()
    text = text.replace("㎎", "mg").replace("／", "/").replace("⁄", "/").replace("²", "2")
    text = re.sub(r"[\[\]\(\)]", "", text)
    text = re.sub(r"\s+", "", text)
    text = text.rstrip(".,")

    aliases = {
        "mmhg": "mmhg",
        "kpa": "kpa",
        "mg/dl": "mg/dl",
        "mgdl": "mg/dl",
        "mmol/l": "mmol/l",
        "mmoll": "mmol/l",
        "mmol/mol": "mmol/mol",
        "mmolmol": "mmol/mol",
        "%": "%",
        "cm": "cm",
        "mm": "mm",
        "m": "m",
        "inch": "in",
        "inches": "in",
        "in": "in",
        '"': "in",
        "kg": "kg",
        "kgs": "kg",
        "kilogram": "kg",
        "kilograms": "kg",
        "lb": "lb",
        "lbs": "lb",
        "pound": "lb",
        "pounds": "lb",
        "kg/m2": "kg/m2",
        "kgm2": "kg/m2",
    }
    return aliases.get(text, text)


def _alias_hit(field_name: str, raw_name: Optional[str]) -> Optional[str]:
    if not raw_name:
        return None

    normalized_name = re.sub(r"\s+", " ", raw_name.strip()).lower()
    for alias in FIELD_SPECS[field_name].aliases:
        normalized_alias = re.sub(r"\s+", " ", alias.strip()).lower()
        if normalized_alias == normalized_name:
            return f"{raw_name}→{field_name}"

    return f"{raw_name}→{field_name}"


#i/p("weight_kg", 100, "lb") -> o/p(45.36, "lb×0.4536", None)
def _normalize_numeric_value(
    field_name: str,
    raw_value: Any,
    raw_unit: Optional[str],
) -> Tuple[Optional[float], Optional[str], Optional[str]]:
    """
    Return:
        processed value,
        conversion description,
        validation reason
    """
    numeric = _coerce_number(raw_value)
    if numeric is None:
        return None, None, "missing_or_non_numeric"

    unit = _normalize_unit_text(raw_unit)
    #Prepare the conversion message: This variable remembers whether a conversion happened.
    factor_description: Optional[str] = None

    if field_name in {"systolic_bp", "diastolic_bp"}:
        if unit == "kpa":
            numeric *= 7.50
            factor_description = "kPa×7.50"
        elif unit in {"", "mmhg"}:
            factor_description = "assumed mmHg" if not unit else None
        else:
            return None, None, f"unsupported_unit:{raw_unit}"

    elif field_name == "blood_sugar":
        if unit == "mg/dl":
            numeric /= 18.0
            factor_description = "mg/dL÷18.0"
        elif unit in {"", "mmol/l"}:
            factor_description = "assumed mmol/L" if not unit else None
        else:
            return None, None, f"unsupported_unit:{raw_unit}"

    elif field_name in {"cholesterol", "ldl", "hdl"}:
        if unit == "mg/dl":
            numeric /= 38.67
            factor_description = "mg/dL÷38.67"
        elif unit in {"", "mmol/l"}:
            factor_description = "assumed mmol/L" if not unit else None
        else:
            return None, None, f"unsupported_unit:{raw_unit}"

    elif field_name == "triglycerides":
        if unit == "mg/dl":
            numeric /= 88.57
            factor_description = "mg/dL÷88.57"
        elif unit in {"", "mmol/l"}:
            factor_description = "assumed mmol/L" if not unit else None
        else:
            return None, None, f"unsupported_unit:{raw_unit}"

    elif field_name == "hba1c":
        if unit == "mmol/mol":
            numeric = numeric * 0.0915 + 2.15
            factor_description = "mmol/mol×0.0915+2.15"
        elif unit in {"", "%"}:
            factor_description = "assumed %" if not unit else None
        else:
            return None, None, f"unsupported_unit:{raw_unit}"

    elif field_name == "height_cm":
        if unit == "in":
            numeric *= 2.54
            factor_description = "inch×2.54"
        elif unit == "m":
            numeric *= 100.0
            factor_description = "m×100"
        elif unit == "mm":
            numeric /= 10.0
            factor_description = "mm÷10"
        elif unit in {"", "cm"}:
            factor_description = "assumed cm" if not unit else None
        else:
            return None, None, f"unsupported_unit:{raw_unit}"

    elif field_name == "weight_kg":
        if unit == "lb":
            numeric *= 0.4536
            factor_description = "lb×0.4536"
        elif unit in {"", "kg"}:
            factor_description = "assumed kg" if not unit else None
        else:
            return None, None, f"unsupported_unit:{raw_unit}"

    elif field_name == "bmi":
        if unit not in {"", "kg/m2"}:
            return None, None, f"unsupported_unit:{raw_unit}"
        factor_description = "assumed kg/m²" if not unit else None

    spec = FIELD_SPECS[field_name]
    if numeric < spec.minimum or numeric > spec.maximum:
        return None, factor_description, (
            f"out_of_range:{numeric:.4g};"
            f"expected={spec.minimum}-{spec.maximum}"
        )

    if spec.value_type == "int":
        return int(round(numeric)), factor_description, None

    return round(float(numeric), 2), factor_description, None

#The function first prepares an empty result box for every expected health value.
def _empty_field_payload(field_name: str) -> Dict[str, Any]:
    return {
        "raw": _empty_raw_item(),
        "processed": {
            "value": None,
            "unit": FIELD_SPECS[field_name].standard_unit,
        },
        "transform": {
            "alias_hit": None,
            "unit_factor": None,
            "reason": "missing",
        },
    }


def _validate_health_values(
    raw_values: Mapping[str, Mapping[str, Any]],
) -> Dict[str, Dict[str, Any]]:
    """Normalize units, convert types, and clamp invalid values to null."""
    #Create empty results for every health field
    fields = {
        field_name: _empty_field_payload(field_name)
        for field_name in FIELD_NAMES
    }

    for field_name in FIELD_NAMES:
        raw_item = raw_values.get(field_name) or {}
        raw_name = raw_item.get("name")
        raw_value = raw_item.get("value")
        raw_unit = raw_item.get("unit")

        processed, unit_factor, reason = _normalize_numeric_value(
            field_name,
            raw_value,
            raw_unit,
        )

        fields[field_name] = {
            "raw": {
                "name": raw_name,
                "value": raw_value,
                "unit": raw_unit,
            },
            "processed": {
                "value": processed,
                "unit": FIELD_SPECS[field_name].standard_unit,
            },
            "transform": {
                "alias_hit": _alias_hit(field_name, raw_name),
                "unit_factor": unit_factor,
                "reason": reason,
            },
        }

    # Derive BMI only when no valid direct BMI is available.
    if fields["bmi"]["processed"]["value"] is None:
        height_cm = fields["height_cm"]["processed"]["value"]
        weight_kg = fields["weight_kg"]["processed"]["value"]

        if height_cm is not None and weight_kg is not None:
            derived_bmi = weight_kg / ((height_cm / 100.0) ** 2)
            spec = FIELD_SPECS["bmi"]

            if spec.minimum <= derived_bmi <= spec.maximum:
                fields["bmi"] = {
                    "raw": {
                        "name": None,
                        "value": None,
                        "unit": None,
                    },
                    "processed": {
                        "value": round(derived_bmi, 2),
                        "unit": spec.standard_unit,
                    },
                    "transform": {
                        "alias_hit": "height_cm+weight_kg→bmi",
                        "unit_factor": "weight_kg÷(height_cm/100)²",
                        "reason": "derived_because_bmi_missing",
                    },
                }

    return fields


def _response_from_fields(
    *,
    status: str,
    message: Optional[str],
    fields: Mapping[str, Mapping[str, Any]],
    additional_fields: Optional[
        list[Dict[str, Any]]
    ] = None,
    include_raw_text: bool = False,
    raw_text: str = "",
) -> Dict[str, Any]:
    response: Dict[str, Any] = {
        "status": status,
        "message": message,
    }

    # Fixed flat keys required by the design's real-time API contract.
    for field_name in FIELD_NAMES:
        response[field_name] = (
            fields.get(field_name, {})
            .get("processed", {})
            .get("value")
        )

    # Dual-layer structure used for user confirmation and persistence.
    response["fields"] = {
        field_name: fields.get(field_name, _empty_field_payload(field_name))
        for field_name in FIELD_NAMES
    }

    response["additional_fields"] = list(
        additional_fields or []
    )

    if include_raw_text:
        response["raw_text"] = raw_text

    return response


def _empty_response(
    status: str,
    message: str,
    *,
    include_raw_text: bool = False,
    raw_text: str = "",
) -> Dict[str, Any]:
    fields = {
        field_name: _empty_field_payload(field_name)
        for field_name in FIELD_NAMES
    }
    return _response_from_fields(
        status=status,
        message=message,
        fields=fields,
        include_raw_text=include_raw_text,
        raw_text=raw_text,
    )


def process_health_report(
    file_bytes: bytes,
    *,
    filename: str = "upload.jpg",
    mime_type: Optional[str] = None,
    include_raw_text: bool = False,
    max_retries: int = 3,
) -> Dict[str, Any]:
    """
    Complete OCR pipeline used by /ocr-health-report.

    no_text:
        No readable transcription.
    no_fields:
        Text exists, but all 11 processed values are null.
    ok:
        At least one processed value is non-null.
    """
    transcription = transcribe_document(
        file_bytes,
        filename=filename,
        mime_type=mime_type,
        max_retries=max_retries,
    )

    if not transcription.strip():
        return _empty_response(
            "no_text",
            "No text was recognized. Please use a clearer image or enter data manually.",
            include_raw_text=include_raw_text,
        )

    fallback_raw = _extract_health_values_regex(transcription)

    additional_fields: list[Dict[str, Any]] = []

    try:
        model_raw, additional_fields = _extract_health_values_from_text(
            transcription,
            max_retries=max_retries,
        )
        raw_values = _merge_raw_extractions(model_raw, fallback_raw)
    except (HealthOcrApiError, json.JSONDecodeError, TypeError, ValueError
            )as error:
                print(
        "❌ Structured OCR extraction failed: "
        f"{type(error).__name__}: {error}"
            )
        # The deterministic fallback keeps the endpoint useful when the
        # text-to-JSON extraction call fails.
                raw_values = fallback_raw

    fields = _validate_health_values(raw_values)
    
    has_canonical_value = any(
        fields[field_name]["processed"]["value"] is not None
        for field_name in FIELD_NAMES
    )
    #the report is considered successful when it contains either one of 11 fields or 1 of additional fields
    has_any_value = (
        has_canonical_value
        or bool(additional_fields)
    )

    if not has_any_value:
        return _response_from_fields(
            status="no_fields",
            message=(
                "This report does not contain recognizable health metrics. "
                "You may add them manually."
            ),
            fields=fields,
            additional_fields=additional_fields,
            include_raw_text=include_raw_text,
            raw_text=transcription,
        )

    return _response_from_fields(
        status="ok",
        message=None,
        fields=fields,
        additional_fields=additional_fields,
        include_raw_text=include_raw_text,
        raw_text=transcription,
    )